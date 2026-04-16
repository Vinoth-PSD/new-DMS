from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

root = Path(r'd:\dms_project')
out_dir = root / 'docs'
out_dir.mkdir(exist_ok=True)
out_path = out_dir / 'DMS_User_Manual_A4.docx'

img = Path(r'C:\Users\vinot\.cursor\projects\d-dms-project\assets\c__Users_vinot_AppData_Roaming_Cursor_User_workspaceStorage_680caac38a52b85f3cbd0ababa6025b8_images_image-6e7c76a8-9a35-4f7a-b292-f8ae2dab11c5.png')

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
normal.font.size = Pt(11)

t = doc.add_paragraph('Document Management System (DMS)')
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].bold = True
t.runs[0].font.size = Pt(20)
sub = doc.add_paragraph('User Manual and EXE Installation Guide')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True

doc.add_paragraph('')

doc.add_heading('1. Purpose', level=1)
doc.add_paragraph('This guide helps non-technical users install the tray EXE and use the web application for daily document processing.')

doc.add_heading('2. Installation (EXE)', level=1)
for s in [
    'Create folder: C:\\DocPro\\Tray',
    'Copy DocProTray.exe and config.json into that folder.',
    'Open config.json and update base_url and folder paths.',
    'Double-click DocProTray.exe. Keep tray icon running.'
]:
    p = doc.add_paragraph(s)
    p.style = doc.styles['List Number']

doc.add_paragraph('Recommended config.json:')
cfg = doc.add_paragraph('{\n  "base_url": "http://YOUR_SERVER:8000",\n  "download_folder": "C:\\\\DocPro\\\\downloads",\n  "upload_folder": "C:\\\\DocPro\\\\upload",\n  "poll_seconds": 5,\n  "watch_seconds": 2,\n  "abbyy_exe_path": "",\n  "process_triggered_open_only": true,\n  "admin_automation_enabled": true,\n  "admin_automation_key": "YOUR_SHARED_KEY",\n  "merged_download_folder": "C:\\\\DocPro\\\\merged",\n  "open_merged_in_word": true\n}')
for r in cfg.runs:
    r.font.name = 'Consolas'
    r.font.size = Pt(9)

doc.add_heading('3. Daily Admin Workflow', level=1)
for s in [
    'Login to admin panel and upload document.',
    'Assign resources and wait for completion.',
    'Click Merge, then Download merged file.',
    'Merged file opens in Word automatically.',
    'Edit and Save. EXE auto-uploads corrected file.',
    'Verify version chips in UI: v1, v2, v3.'
]:
    p = doc.add_paragraph(s)
    p.style = doc.styles['List Number']

doc.add_heading('4. Troubleshooting', level=1)
for s in [
    'EXE does not start: ensure config.json is beside EXE.',
    '403 upload error: admin_automation_key must match server ADMIN_AUTOMATION_KEY.',
    'No auto-upload after save: save inside merged_download_folder and keep tray running.',
    'Download issue: verify base_url and backend service status.'
]:
    p = doc.add_paragraph(s)
    p.style = doc.styles['List Bullet']

doc.add_heading('5. Screenshot Reference', level=1)
if img.exists():
    doc.add_paragraph('Screenshot: config.json setup example')
    doc.add_picture(str(img), width=Inches(6.7))
else:
    doc.add_paragraph('Screenshot file not found during generation.')

doc.save(out_path)
print(out_path)
